"""Regression tests for deterministic JD-to-resume catalog selection."""

import json
from pathlib import Path
import re
import sys
import unittest

from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from fit import detect_family
from resume_engine import generate, tailor_for_job
from resume_engine.catalog import load_catalog


ROOT = Path(__file__).parent
FIXTURES = json.loads(
    (ROOT / "fixtures" / "representative_jds.json").read_text(encoding="utf-8")
)
OUTPUT = ROOT.parent / "output" / "catalog_test"
UNSUPPORTED_SENIORITY = re.compile(
    r"\b(senior|sr\.?|lead|manager|director|principal|head|chief|vp)\b", re.I
)


class CatalogTailoringTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        OUTPUT.mkdir(parents=True, exist_ok=True)

    def test_representative_jds(self):
        for fixture in FIXTURES:
            with self.subTest(fixture=fixture["name"]):
                family, _ = detect_family(fixture["title"], fixture["jd_text"])
                self.assertEqual(fixture["expected_family"], family)

                data, report = tailor_for_job(
                    family, fixture["title"], fixture["jd_text"]
                )
                self.assertEqual(5, len(data.experiences))
                self.assertEqual("Jan 2022 – Aug 2022", data.experiences[-1].date)
                self.assertLessEqual(
                    sum("intern" in exp.role.casefold() for exp in data.experiences),
                    2,
                )
                for exp in data.experiences:
                    self.assertIsNone(UNSUPPORTED_SENIORITY.search(exp.role), exp.role)
                self.assertEqual([4, 3, 4, 3, 1], [len(exp.bullets) for exp in data.experiences])

                skill_text = " | ".join(category.skills for category in data.skills)
                for expected in fixture["expected_skills"]:
                    self.assertIn(expected, skill_text)
                for expected in fixture["expected_gaps"]:
                    self.assertIn(expected, report["unapproved_jd_terms"])

                self.assertLessEqual(len(data.skills), 4)
                for category in data.skills:
                    self.assertLessEqual(len(category.skills), 160)

                out = OUTPUT / f"{fixture['name']}.docx"
                _, warnings = generate(data, str(out))
                self.assertEqual([], warnings)
                doc = Document(out)
                self.assertEqual(1, len(doc.tables))
                self.assertEqual(21, len(doc.tables[0].rows))

    def test_catalog_safety(self):
        catalog = load_catalog()
        categories = {item["name"] for item in catalog["skill_categories"]}
        companies = set()
        for experience in catalog["experiences"]:
            self.assertNotIn(experience["company"], companies)
            companies.add(experience["company"])
            for title in experience["titles"]:
                self.assertIsNone(UNSUPPORTED_SENIORITY.search(title["text"]), title["text"])
            for bullet in experience["bullets"]:
                self.assertLessEqual(len(bullet["text"]), 340, bullet["text"])

        skill_names = set()
        for skill in catalog["skills"]:
            self.assertIn(skill["category"], categories)
            self.assertNotIn(skill["name"].casefold(), skill_names)
            skill_names.add(skill["name"].casefold())

    def test_business_analyst_profile_matches_approved_reference(self):
        data, _ = tailor_for_job(
            "Business Analyst",
            "Business Analyst",
            "Requirements gathering, user stories, acceptance criteria, UAT, "
            "Agile/Scrum, process mapping, SQL, Power BI, and stakeholder management.",
        )

        self.assertEqual(
            [3, 4, 5, 4, 2],
            [len(experience.bullets) for experience in data.experiences],
        )
        self.assertEqual("Jan 2022 – Aug 2022", data.experiences[-1].date)
        self.assertEqual(
            ["Business Analysis", "Tools & Analytics", "Strategy & Operations"],
            [category.name for category in data.skills],
        )
        bullet_text = "\n".join(
            bullet
            for experience in data.experiences
            for bullet in experience.bullets
        )
        for phrase in (
            "requirements gathering between business stakeholders and UI/UX designers",
            "user stories and acceptance criteria",
            "Supported UAT by validating reports",
            "documentation of business requirements for Client Revenue Operations",
            "healthcare business datasets using SQL, Python, and Excel",
        ):
            self.assertIn(phrase, bullet_text)

        out = OUTPUT / "business_analyst_reference.docx"
        _, warnings = generate(data, str(out))
        self.assertEqual([], warnings)
        table = Document(out).tables[0]
        self.assertEqual(
            [
                "PROFESSIONAL SUMMARY",
                "EXPERIENCE",
                "EDUCATION",
                "SKILLS",
                "HONORS & AWARDS",
            ],
            [table.rows[index].cells[0].text for index in (2, 4, 15, 17, 19)],
        )
        education = table.rows[16].cells[0].paragraphs
        self.assertEqual(2, len(education))
        self.assertIn("GPA: 3.8\tJune 2026", education[0].text)
        self.assertIn("GPA: 3.5\tAug 2023", education[1].text)
        self.assertIn(
            "Process Mapping, SOPs, Root-Cause Analysis, BRD, FRD",
            table.rows[18].cells[0].text,
        )


if __name__ == "__main__":
    unittest.main()
