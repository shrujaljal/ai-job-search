"""
Resume generator that produces Word documents matching Shrujal's resume template.

Layout: single borderless table, A4 page, 0.5" margins, 10pt Calibri.
Three-column grid: 7387 + 184 + 2895 twips.
Section order: summary, experience, education, skills, honors/awards.
Section headers are bold blue (#2F5496) with a bottom rule.
"""

from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import ResumeData
from .content_rules import enforce

# ── Layout constants (twips; 1440 twips = 1 inch) ───────────────────────────
COL1_W = 7387
COL2_W = 184
COL3_W = 2895
LEFT_W = COL1_W + COL2_W
TOTAL_W = COL1_W + COL2_W + COL3_W

# ── Style constants ──────────────────────────────────────────────────────────
BLUE = "2F5496"
FONT = "Calibri"
SZ = 20        # half-points → 10pt
SZ_NAME = 32   # half-points → 16pt

TEMPLATE = Path(__file__).parent / "new_template.docx"
if not TEMPLATE.exists():
    TEMPLATE = Path(__file__).parent / "base_template.docx"


# ════════════════════════════════════════════════════════════════════════════
# Low-level XML helpers
# ════════════════════════════════════════════════════════════════════════════

def _rpr(bold=False, italic=False, color=None, sz=SZ) -> OxmlElement:
    """Build a <w:rPr> element."""
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:cstheme"), "minorHAnsi")
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if italic:
        rPr.append(OxmlElement("w:i"))
        rPr.append(OxmlElement("w:iCs"))
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        if color == BLUE:
            c.set(qn("w:themeColor"), "accent1")
            c.set(qn("w:themeShade"), "BF")
        rPr.append(c)
    s = OxmlElement("w:sz")
    s.set(qn("w:val"), str(sz))
    rPr.append(s)
    sCs = OxmlElement("w:szCs")
    sCs.set(qn("w:val"), str(sz))
    rPr.append(sCs)
    return rPr


def _run(text: str, bold=False, italic=False, color=None, sz=SZ) -> OxmlElement:
    """Build a <w:r> element with text."""
    r = OxmlElement("w:r")
    r.append(_rpr(bold=bold, italic=italic, color=color, sz=sz))
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _hyperlink_run(rel_id: str, display: str) -> OxmlElement:
    """Build a <w:hyperlink> element."""
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rel_id)
    hl.set(qn("w:history"), "1")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:cstheme"), "minorHAnsi")
    rPr.append(style)
    rPr.append(fonts)
    s = OxmlElement("w:sz")
    s.set(qn("w:val"), str(SZ))
    rPr.append(s)
    sCs = OxmlElement("w:szCs")
    sCs.set(qn("w:val"), str(SZ))
    rPr.append(sCs)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = display
    r.append(t)
    hl.append(r)
    return hl


def _line_break() -> OxmlElement:
    return OxmlElement("w:br")


def _tab() -> OxmlElement:
    run = OxmlElement("w:r")
    run.append(_rpr(sz=SZ))
    run.append(OxmlElement("w:tab"))
    return run


def _add_right_tab_stop(paragraph: OxmlElement, position=10080) -> None:
    p_pr = paragraph.find(qn("w:pPr"))
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position))
    tabs.append(tab)
    p_pr.append(tabs)


def _external_rel_id(doc: Document, url: str) -> str:
    return doc.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)


def _para(align=None, space_before=None, space_after=None, justify=False) -> OxmlElement:
    """Build a <w:p> element with optional paragraph properties."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    if justify:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)
    elif align:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), align)
        pPr.append(jc)
    if space_before is not None or space_after is not None:
        sp = OxmlElement("w:spacing")
        if space_before is not None:
            sp.set(qn("w:before"), str(space_before))
        if space_after is not None:
            sp.set(qn("w:after"), str(space_after))
        pPr.append(sp)
    p.append(pPr)
    return p


def _bullet_para(num_id: str, text: str) -> OxmlElement:
    """Build a bullet-list paragraph."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "ListParagraph")
    pPr.append(style)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    rpr_el = _rpr(sz=SZ)
    pPr.append(rpr_el)
    p.append(pPr)
    p.append(_run(text, sz=SZ))
    return p


def _new_row(table) -> OxmlElement:
    """Add a blank row to the table and return its <w:tr> element."""
    tr = OxmlElement("w:tr")
    table._tbl.append(tr)
    return tr


def _span_cell(tr, content_paras: list, bottom_border=False) -> None:
    """Add a single cell spanning the full three-column table."""
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(TOTAL_W))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    gridSpan = OxmlElement("w:gridSpan")
    gridSpan.set(qn("w:val"), "3")
    tcPr.append(gridSpan)
    if bottom_border:
        tcBorders = OxmlElement("w:tcBorders")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "0")
        bot.set(qn("w:color"), "auto")
        tcBorders.append(bot)
        tcPr.append(tcBorders)
    tc.append(tcPr)
    for p in content_paras:
        tc.append(p)
    tr.append(tc)


def _two_cells(tr, left_paras: list, right_paras: list,
               left_w=LEFT_W, right_w=COL3_W) -> None:
    """Add two cells to a row."""
    for index, (w, paras) in enumerate(
        [(left_w, left_paras), (right_w, right_paras)]
    ):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(w))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)
        if index == 0:
            gridSpan = OxmlElement("w:gridSpan")
            gridSpan.set(qn("w:val"), "2")
            tcPr.append(gridSpan)
        tc.append(tcPr)
        for p in paras:
            tc.append(p)
        tr.append(tc)


def _set_table_geometry(table) -> None:
    """Keep the table width, grid, and generated cell widths in agreement."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.insert(0, tblW)
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is not None:
        tblPr.remove(layout)

    grid = tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for width in (COL1_W, COL2_W, COL3_W):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


# ════════════════════════════════════════════════════════════════════════════
# Section builders
# ════════════════════════════════════════════════════════════════════════════

def _build_header_row(table, data: ResumeData, num_id: str) -> None:
    tr = _new_row(table)
    p = _para(align="center")
    p.append(_run(data.name, bold=False, color=BLUE, sz=SZ_NAME))
    _span_cell(tr, [p])


def _build_contact_row(table, data: ResumeData, rel_ids: dict[str, str]) -> None:
    tr = _new_row(table)
    p = _para(align="center", space_before=60)
    p.append(_run(f"{data.phone} | ", sz=SZ))
    p.append(_hyperlink_run(rel_ids["email"], data.email))
    p.append(_run(" | ", sz=SZ))
    p.append(_hyperlink_run(rel_ids["linkedin"], data.linkedin_text))
    if data.portfolio_url:
        p.append(_run(" | ", sz=SZ))
        p.append(_hyperlink_run(rel_ids["portfolio"], data.portfolio_text))
    if data.github_url:
        p.append(_run(" | ", sz=SZ))
        p.append(_hyperlink_run(rel_ids["github"], data.github_text))
    _span_cell(tr, [p])


def _build_section_header(table, label: str, space_before=120) -> None:
    tr = _new_row(table)
    p = _para(space_before=space_before)
    p.find(qn("w:pPr")).append(_rpr(bold=True, color=BLUE, sz=SZ))
    p.append(_run(label, bold=True, color=BLUE, sz=SZ))
    _span_cell(tr, [p], bottom_border=True)


def _build_summary_row(table, data: ResumeData) -> None:
    tr = _new_row(table)
    p = _para(justify=True)
    p.append(_run(data.summary, sz=SZ))
    _span_cell(tr, [p])


def _build_experience_header(
    table, company: str, role: str, date: str, space_before=60
) -> None:
    tr = _new_row(table)
    # Left: "Company | Role"
    lp = _para(justify=True, space_before=space_before)
    lp.append(_run(f"{company} | ", bold=True, sz=SZ))
    lp.append(_run(role, bold=True, sz=SZ))
    # Right: date, right-aligned
    rp = _para(align="right", space_before=space_before)
    rp.append(_run(date, bold=True, sz=SZ))
    _two_cells(tr, [lp], [rp])


def _build_bullets_row(table, bullets: list[str], num_id: str) -> None:
    tr = _new_row(table)
    paras = [_bullet_para(num_id, b) for b in bullets]
    _span_cell(tr, paras)


def _build_education_row(table, data: ResumeData) -> None:
    tr = _new_row(table)
    paras = []

    # MBA line
    p1 = _para()
    p1.append(_run("Master of Business Administration (MBA)", bold=True, sz=SZ))
    p1.append(_run(
        " - University of California, Riverside (GPA: 3.8) — June 2026", sz=SZ
    ))
    paras.append(p1)

    # Honors
    p2 = _para()
    p2.append(_run("Honors: Beta Gamma Sigma Award, 2024 Case Competition Winner",
                   italic=True, sz=SZ))
    paras.append(p2)

    # Coursework + B.Tech on same paragraph (line break between)
    p3 = _para()
    p3.append(_run("Relevant Coursework: ", italic=True, sz=SZ))
    p3.append(_run(data.coursework, italic=True, sz=SZ))
    br = OxmlElement("w:br")
    p3.append(br)
    p3.append(_run("Bachelor of Technology, Biotechnology ", bold=True, sz=SZ))
    p3.append(_run(
        "– Vellore Institute of Technology, Vellore, India – Aug 2023", sz=SZ
    ))
    paras.append(p3)

    _span_cell(tr, paras)


def _build_projects_row(table, data: ResumeData, num_id: str) -> None:
    tr = _new_row(table)
    paras = []

    for proj in data.projects:
        # Project title
        tp = _para()
        tp.append(_run(proj.title, bold=True, sz=SZ))
        paras.append(tp)
        for b in proj.bullets:
            paras.append(_bullet_para(num_id, b))

    # Leadership bullets (inline after projects)
    if data.leadership_bullets:
        # Leadership title line
        lp = _para()
        lp.append(_run("UCR GSM Student Association | Marketing Lead", bold=True, sz=SZ))
        paras.append(lp)
        for b in data.leadership_bullets:
            paras.append(_bullet_para(num_id, b))

    _span_cell(tr, paras)


def _build_skills_row(table, data: ResumeData, num_id: str) -> None:
    tr = _new_row(table)
    paras = []
    for cat in data.skills:
        p = _bullet_para(num_id, "")
        # Remove the plain text run and replace with bold label + normal text
        for child in list(p):
            if child.tag == qn("w:r"):
                p.remove(child)
        r_bold = _run(f"{cat.name}:", bold=True, sz=SZ)
        r_text = _run(f" {cat.skills}", sz=SZ)
        p.append(r_bold)
        p.append(r_text)
        paras.append(p)
    _span_cell(tr, paras)


# Replacement builders for the current DOCX template. They intentionally appear
# after the legacy builders above so these definitions are the ones generate()
# uses at runtime.
def _build_education_row(table, data: ResumeData) -> None:
    tr = _new_row(table)
    paragraphs = []

    p1 = _para()
    _add_right_tab_stop(p1)
    p1.append(_run(
        "Master of Business Administration (STEM MBA) \u2013 University of California, Riverside",
        bold=True, sz=SZ,
    ))
    p1.append(_run(" GPA: 3.8", italic=True, sz=SZ))
    p1.append(_tab())
    p1.append(_run("June 2026", bold=True, sz=SZ))
    paragraphs.append(p1)

    p2 = _para()
    _add_right_tab_stop(p2)
    p2.append(_run(
        "Bachelor of Technology, Biotechnology \u2013 Vellore Institute of Technology, India",
        bold=True, sz=SZ,
    ))
    p2.append(_run(" GPA: 3.5", italic=True, sz=SZ))
    p2.append(_tab())
    p2.append(_run("Aug 2023", bold=True, sz=SZ))
    paragraphs.append(p2)

    _span_cell(tr, paragraphs)


def _build_projects_row(table, data: ResumeData, num_id: str) -> None:
    tr = _new_row(table)
    paras = []

    for proj in data.projects:
        tp = _para()
        tp.append(_run(proj.title, bold=True, sz=SZ))
        paras.append(tp)
        for b in proj.bullets:
            paras.append(_bullet_para(num_id, b))

    if data.leadership_bullets:
        lp = _para()
        lp.append(_run(
            "UCR GSM Student Association | Professional Development Lead",
            bold=True, sz=SZ,
        ))
        paras.append(lp)
        for b in data.leadership_bullets:
            paras.append(_bullet_para(num_id, b))

    _span_cell(tr, paras)


def _build_skills_row(table, data: ResumeData, num_id: str = "27") -> None:
    tr = _new_row(table)
    paras = []
    for cat in data.skills:
        p = _para(justify=True)
        p.append(_run(f"{cat.name}:", bold=True, sz=SZ))
        p.append(_run(f" {cat.skills}", sz=SZ))
        paras.append(p)
    _span_cell(tr, paras)


def _build_honors_row(table, data: ResumeData, num_id: str) -> None:
    tr = _new_row(table)
    paras = []
    for index in range(0, len(data.honors_awards), 2):
        p = _para(justify=True)
        p.append(_run(" | ".join(data.honors_awards[index:index + 2]), sz=SZ))
        paras.append(p)
    _span_cell(tr, paras)


# ════════════════════════════════════════════════════════════════════════════
# Public API
# ════════════════════════════════════════════════════════════════════════════

_SECTION_ROWS = {
    2: "PROFESSIONAL SUMMARY",
    4: "EXPERIENCE",
    15: "EDUCATION",
    17: "SKILLS",
    19: "HONORS & AWARDS",
}
_EXPERIENCE_ROWS = ((5, 6), (7, 8), (9, 10), (11, 12), (13, 14))


def _unique_row_cells(row) -> list:
    cells = []
    seen = set()
    for cell in row.cells:
        marker = id(cell._tc)
        if marker not in seen:
            seen.add(marker)
            cells.append(cell)
    return cells


def _replace_paragraph_element_text(paragraph, text: str) -> None:
    prototype_rpr = None
    for run in paragraph.iter(qn("w:r")):
        run_rpr = run.find(qn("w:rPr"))
        if run_rpr is not None:
            prototype_rpr = deepcopy(run_rpr)
            break

    for child in list(paragraph):
        if child.tag != qn("w:pPr"):
            paragraph.remove(child)

    run = OxmlElement("w:r")
    run.append(prototype_rpr if prototype_rpr is not None else _rpr(sz=SZ))
    text_element = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(
            "{http://www.w3.org/XML/1998/namespace}space", "preserve"
        )
    text_element.text = text
    run.append(text_element)
    paragraph.append(run)


def _replace_paragraph_text(paragraph, text: str) -> None:
    _replace_paragraph_element_text(paragraph._p, text)


def _replace_cell_paragraphs(cell, texts: list[str]) -> None:
    prototypes = [deepcopy(paragraph._p) for paragraph in cell.paragraphs]
    if not prototypes:
        prototypes = [_para()]

    for paragraph in list(cell._tc.findall(qn("w:p"))):
        cell._tc.remove(paragraph)

    for index, text in enumerate(texts):
        paragraph = deepcopy(prototypes[min(index, len(prototypes) - 1)])
        _replace_paragraph_element_text(paragraph, text)
        cell._tc.append(paragraph)

    if not texts:
        paragraph = deepcopy(prototypes[0])
        _replace_paragraph_element_text(paragraph, "")
        cell._tc.append(paragraph)


def _replace_skill_paragraph(paragraph, name: str, skills: str) -> None:
    run_properties = []
    for run in paragraph.iter(qn("w:r")):
        run_rpr = run.find(qn("w:rPr"))
        if run_rpr is not None:
            run_properties.append(deepcopy(run_rpr))

    for child in list(paragraph):
        if child.tag != qn("w:pPr"):
            paragraph.remove(child)

    for index, text in enumerate((f"{name}:", f" {skills}")):
        run = OxmlElement("w:r")
        if run_properties:
            run.append(deepcopy(run_properties[min(index, len(run_properties) - 1)]))
        else:
            run.append(_rpr(bold=index == 0, sz=SZ))
        text_element = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            text_element.set(
                "{http://www.w3.org/XML/1998/namespace}space", "preserve"
            )
        text_element.text = text
        run.append(text_element)
        paragraph.append(run)


def _replace_skills(cell, data: ResumeData) -> None:
    prototypes = [deepcopy(paragraph._p) for paragraph in cell.paragraphs]
    if not prototypes:
        prototypes = [_para(justify=True)]

    for paragraph in list(cell._tc.findall(qn("w:p"))):
        cell._tc.remove(paragraph)

    for index, category in enumerate(data.skills):
        paragraph = deepcopy(prototypes[min(index, len(prototypes) - 1)])
        _replace_skill_paragraph(paragraph, category.name, category.skills)
        cell._tc.append(paragraph)


def _template_whitespace(text: str) -> tuple[str, str]:
    return text[:len(text) - len(text.lstrip())], text[len(text.rstrip()):]


def _validate_template(table) -> None:
    if len(table.rows) != 21:
        raise ValueError(
            f"Resume template must contain 21 rows; found {len(table.rows)}."
        )
    for row_index, expected in _SECTION_ROWS.items():
        actual = table.rows[row_index].cells[0].text.strip()
        if actual != expected:
            raise ValueError(
                f"Resume template row {row_index} must be {expected!r}; "
                f"found {actual!r}."
            )


def _fill_template(table, data: ResumeData) -> None:
    """Fill the authoritative template without rebuilding its formatting."""
    _validate_template(table)

    _replace_paragraph_text(table.rows[0].cells[0].paragraphs[0], data.name)
    _replace_paragraph_text(table.rows[3].cells[0].paragraphs[0], data.summary)

    for experience, (header_row, bullet_row) in zip(
        data.experiences, _EXPERIENCE_ROWS
    ):
        header_cells = _unique_row_cells(table.rows[header_row])
        if len(header_cells) != 2:
            raise ValueError(
                f"Resume template row {header_row} must contain two unique cells."
            )

        title_paragraph = header_cells[0].paragraphs[0]
        date_paragraph = header_cells[1].paragraphs[0]
        _, title_suffix = _template_whitespace(title_paragraph.text)
        date_prefix, date_suffix = _template_whitespace(date_paragraph.text)
        _replace_paragraph_text(
            title_paragraph,
            f"{experience.company} | {experience.role}{title_suffix}",
        )
        _replace_paragraph_text(
            date_paragraph,
            f"{date_prefix}{experience.date}{date_suffix}",
        )

        bullet_cell = _unique_row_cells(table.rows[bullet_row])[0]
        _replace_cell_paragraphs(bullet_cell, experience.bullets)

    skills_cell = _unique_row_cells(table.rows[18])[0]
    _replace_skills(skills_cell, data)


def generate(data: ResumeData, output_path: str) -> tuple[str, list[str]]:
    """
    Generate a .docx resume from ResumeData.

    Returns (output_path, warnings).
    Enforces 1-page content rules before generating.
    """
    data, warnings = enforce(data)

    doc = Document(str(TEMPLATE))

    table = doc.tables[0]
    _fill_template(table, data)

    doc.save(output_path)
    return output_path, warnings
