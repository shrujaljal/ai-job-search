from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

import config
import main
import profile_import
import resume_render


MASTER_MARKDOWN = """# Jane Candidate
jane@example.com | linkedin.com/in/jane-candidate | California

# Professional Experience
## Acme | Operations Analyst | 2024 - Present
- Built weekly Excel dashboards for leadership reporting.
- Documented reporting workflows across operations teams.

# Education
## MBA | Example University | 2026
- Beta Gamma Sigma Honor Society

# Honors & Awards
- Case Competition Winner

# Technical Skills
Tools: Excel, SQL, Tableau
"""


class ProfileImportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.original_data_dir = config.DATA_DIR
        config.DATA_DIR = Path(self.temp.name) / "data"
        config.ensure_config()

    def tearDown(self) -> None:
        config.DATA_DIR = self.original_data_dir
        self.temp.cleanup()

    def test_markdown_import_extracts_identity_order_and_sections(self) -> None:
        path = Path(self.temp.name) / "master.md"
        path.write_text(MASTER_MARKDOWN, encoding="utf-8")

        parsed = profile_import.parse_resume(profile_import.extract_lines(path), path.name)

        self.assertEqual(parsed["identity"]["name"], "Jane Candidate")
        self.assertEqual(parsed["identity"]["email"], "jane@example.com")
        self.assertEqual(parsed["experience"][0]["company"], "Acme")
        self.assertEqual(parsed["experience"][0]["role"], "Operations Analyst")
        self.assertEqual(parsed["honors"], ["Case Competition Winner"])
        self.assertEqual(
            parsed["resume_blueprint"]["section_order"],
            ["experience", "education", "honors", "skills"],
        )
        self.assertEqual(parsed["resume_blueprint"]["section_headings"]["experience"], "Professional Experience")

    def test_multiple_files_merge_unique_bullets_into_matching_experience(self) -> None:
        first = Path(self.temp.name) / "master.md"
        second = Path(self.temp.name) / "expanded.md"
        first.write_text(MASTER_MARKDOWN, encoding="utf-8")
        second.write_text("""# Experience
## Acme | Operations Analyst | 2024 - Present
- Built weekly Excel dashboards for leadership reporting.
- Built weekly Excel dashboards used for leadership reporting.
- Coordinated monthly KPI reviews with cross-functional partners.
""", encoding="utf-8")
        documents = [
            profile_import.parse_resume(profile_import.extract_lines(path), path.name)
            for path in (first, second)
        ]

        merged, stats = profile_import.merge_profile(config.load("profile"), documents)

        bullets = merged["experience"][0]["bullets"]
        self.assertEqual(len(bullets), 3)
        self.assertIn("Coordinated monthly KPI reviews with cross-functional partners.", bullets)
        self.assertGreaterEqual(stats["duplicates_removed"], 1)
        self.assertNotIn("Built weekly Excel dashboards used for leadership reporting.", bullets)
        self.assertEqual(merged["resume_blueprint"]["section_order"][0], "experience")

    def test_same_role_with_different_dates_remains_separate(self) -> None:
        existing = {
            "experience": [{
                "company": "Acme", "role": "Analyst", "date": "2022",
                "bullets": ["Supported the first team."],
            }],
        }
        incoming = {
            "experience": [{
                "company": "Acme", "role": "Analyst", "date": "2024",
                "bullets": ["Supported the second team."],
            }],
        }

        merged, _ = profile_import.merge_profile(existing, [incoming])

        self.assertEqual(len(merged["experience"]), 2)

    def test_docx_lists_and_headings_are_imported(self) -> None:
        path = Path(self.temp.name) / "master.docx"
        document = Document()
        document.add_paragraph("Jane Candidate", style="Title")
        document.add_paragraph("jane@example.com")
        document.add_heading("Experience", level=1)
        document.add_heading("Acme | Analyst | 2024", level=2)
        document.add_paragraph("Created an operational dashboard.", style="List Bullet")
        document.add_heading("Skills", level=1)
        document.add_paragraph("Tools: Excel, SQL")
        document.save(path)

        parsed = profile_import.parse_resume(profile_import.extract_lines(path), path.name)

        self.assertEqual(parsed["experience"][0]["bullets"], ["Created an operational dashboard."])
        self.assertEqual(parsed["skills"], [{"name": "Tools", "items": "Excel, SQL"}])

    def test_pdf_text_is_accepted(self) -> None:
        path = Path(self.temp.name) / "master.pdf"
        writer = PdfWriter()
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        font_ref = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject({
            NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref}),
        })
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 14 TL 72 720 Td (Jane Candidate) Tj T* (EXPERIENCE) Tj T* (- Built dashboards) Tj ET")
        page[NameObject("/Contents")] = writer._add_object(stream)
        with path.open("wb") as output:
            writer.write(output)

        lines = profile_import.extract_lines(path)

        self.assertTrue(any(line.text == "Jane Candidate" for line in lines))
        self.assertTrue(any(line.text == "EXPERIENCE" and line.heading for line in lines))

    def test_renderer_follows_blueprint_order_and_headings(self) -> None:
        path = Path(self.temp.name) / "resume.docx"
        context = {
            "identity": {"name": "Jane Candidate", "email": "jane@example.com"},
            "summary": "Operations analyst.",
            "experiences": [{"company": "Acme", "role": "Analyst", "date": "2024", "bullets": ["Built dashboards."]}],
            "education": [], "projects": [], "leadership": [],
            "skills": [{"name": "Tools", "items": "Excel"}],
            "honors": [], "custom_sections": [], "coursework": "",
            "resume_blueprint": {
                "section_order": ["skills", "summary", "experience"],
                "section_headings": {"skills": "Capabilities", "summary": "Profile", "experience": "Career History"},
            },
        }

        resume_render.render_docx(context, str(path))
        text = "\n".join(cell.text for table in Document(path).tables for row in table.rows for cell in row.cells)

        self.assertLess(text.index("Capabilities"), text.index("Profile"))
        self.assertLess(text.index("Profile"), text.index("Career History"))

    def test_table_based_resume_round_trips_into_profile(self) -> None:
        path = Path(self.temp.name) / "table-resume.docx"
        context = {
            "identity": {"name": "Jane Candidate", "email": "jane@example.com"},
            "summary": "Operations analyst.",
            "experiences": [{"company": "Acme", "role": "Analyst", "date": "2024", "bullets": ["Built dashboards."]}],
            "education": [], "projects": [], "leadership": [], "skills": [],
            "honors": [], "custom_sections": [], "coursework": "",
            "resume_blueprint": {
                "section_order": ["summary", "experience"],
                "section_headings": {"summary": "Professional Summary", "experience": "Experience"},
            },
        }
        resume_render.render_docx(context, str(path))

        parsed = profile_import.parse_resume(profile_import.extract_lines(path), path.name)

        self.assertEqual(parsed["identity"]["name"], "Jane Candidate")
        self.assertEqual(parsed["experience"][0]["company"], "Acme")
        self.assertEqual(parsed["experience"][0]["bullets"], ["Built dashboards."])

    def test_http_import_persists_sources_and_prompt(self) -> None:
        client = TestClient(main.app)

        response = client.post(
            "/api/profile/import",
            files=[("files", ("master.md", MASTER_MARKDOWN.encode(), "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(payload["stats"]["files"], 1)
        self.assertEqual(payload["sources"], ["master.md"])
        self.assertEqual(
            payload["profile"]["resume_blueprint"]["section_headings"]["experience"],
            "Professional Experience",
        )
        self.assertTrue((config.DATA_DIR / "profile_sources" / "master.md").exists())
        prompt_response = client.get("/api/profile/enrichment-prompt")
        self.assertEqual(prompt_response.status_code, 200)
        self.assertIn("Prompt 1: Acme", prompt_response.text)
        self.assertIn("# Experience", prompt_response.text)


if __name__ == "__main__":
    unittest.main()
