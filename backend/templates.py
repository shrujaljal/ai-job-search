"""User-managed DOCX resume templates.

The shipped default template is rendered by resume_render's native builder. User
templates are DOCX files with placeholder tokens such as {{summary}} and
{{experience}}. This module validates, lists, stores, and resolves those files.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

import config

DEFAULT_TEMPLATE_ID = "default"
REQUIRED_TOKENS = {"summary", "experience", "education", "skills"}
KNOWN_TOKENS = {
    "name",
    "contact",
    "location",
    "phone",
    "email",
    "links",
    "summary",
    "experience",
    "education",
    "coursework",
    "projects",
    "leadership",
    "skills",
    "family",
}
TOKEN_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")
SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9._ -]+")


def templates_dir() -> Path:
    config.ensure_config()
    d = config.DATA_DIR / "templates"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _safe_filename(name: str) -> str:
    stem = Path(name).stem.strip() or "template"
    stem = SAFE_NAME_RE.sub("", stem).strip(" ._-") or "template"
    return f"{stem[:80]}.docx"


def _unique_path(filename: str) -> Path:
    base = templates_dir() / _safe_filename(filename)
    if not base.exists():
        return base
    stem = base.stem
    for i in range(2, 1000):
        candidate = base.with_name(f"{stem}-{i}.docx")
        if not candidate.exists():
            return candidate
    raise ValueError("Could not allocate a unique template filename.")


def _doc_text(path: Path) -> str:
    doc = Document(str(path))
    chunks: list[str] = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


def validate_template(path: Path) -> dict:
    tokens = sorted(set(TOKEN_RE.findall(_doc_text(path))))
    missing = sorted(REQUIRED_TOKENS - set(tokens))
    unknown = sorted(set(tokens) - KNOWN_TOKENS)
    return {
        "valid": not missing,
        "recognized_tokens": tokens,
        "missing_tokens": missing,
        "unknown_tokens": unknown,
    }


def list_templates(active_id: str | None = None) -> list[dict]:
    active = active_id or DEFAULT_TEMPLATE_ID
    items = [{
        "id": DEFAULT_TEMPLATE_ID,
        "name": "Built-in default",
        "builtin": True,
        "active": active == DEFAULT_TEMPLATE_ID,
        "valid": True,
        "recognized_tokens": sorted(KNOWN_TOKENS),
        "missing_tokens": [],
        "unknown_tokens": [],
        "size": 0,
    }]
    for path in sorted(templates_dir().glob("*.docx")):
        info = validate_template(path)
        items.append({
            "id": path.name,
            "name": path.stem,
            "builtin": False,
            "active": active == path.name,
            "size": path.stat().st_size,
            **info,
        })
    return items


def save_upload(filename: str, src_path: Path) -> dict:
    info = validate_template(src_path)
    dst = _unique_path(filename)
    shutil.copyfile(src_path, dst)
    return {"id": dst.name, "name": dst.stem, "builtin": False, "size": dst.stat().st_size, **info}


def delete_template(template_id: str) -> None:
    if template_id == DEFAULT_TEMPLATE_ID:
        raise ValueError("The built-in default template cannot be deleted.")
    path = resolve_template(template_id)
    if not path or not path.exists():
        raise FileNotFoundError(template_id)
    path.unlink()


def resolve_template(template_id: str | None) -> Path | None:
    if not template_id or template_id == DEFAULT_TEMPLATE_ID:
        return None
    candidate = (templates_dir() / Path(template_id).name).resolve()
    root = templates_dir().resolve()
    if candidate != root and root in candidate.parents and candidate.suffix.lower() == ".docx":
        return candidate
    return None
